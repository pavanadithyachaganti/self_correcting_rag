import os

SUPPORTED = (".md", ".txt", ".pdf", ".docx")


def extract_text(path):
    ext = os.path.splitext(path)[1].lower()
    if ext in (".md", ".txt"):
        with open(path, encoding="utf-8", errors="ignore") as f:
            return f.read()
    if ext == ".pdf":
        return _extract_pdf(path)
    if ext == ".docx":
        return _extract_docx(path)
    return ""


def _extract_pdf(path):
    # Prefer PyMuPDF for cleaner multi-column extraction; fall back to pypdf.
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(path)
        parts = [p.get_text("text") for p in doc]
        text = "\n\n".join(t for t in parts if t.strip())
        if text.strip():
            return text
    except Exception:
        pass
    from pypdf import PdfReader
    reader = PdfReader(path)
    parts = []
    for page in reader.pages:
        t = page.extract_text() or ""
        if t.strip():
            parts.append(t)
    return "\n\n".join(parts)


def _extract_docx(path):
    from docx import Document
    doc = Document(path)
    blocks = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                blocks.append(" | ".join(cells))
    return "\n\n".join(blocks)